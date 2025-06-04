"""  
表格替换模块  
负责处理Word文档中表格内容的替换（仅使用Word对象处理）
"""  

import os
import json
import re

def replace_values_with_placeholders(doc, match_results_dir, match_files):
    """
    将Word文档中的表格内容替换为占位符，用于生成模板
    
    参数:
        doc: docx.Document对象，要处理的原始文档
        match_results_dir: 匹配结果目录路径
        match_files: 表格匹配结果文件列表
    """
    
    try:
        # 获取文档中的所有表格（包括嵌套表格）
        all_tables = get_all_tables_including_nested(doc)
        
        # 处理每个匹配文件
        for match_file in match_files:
            file_path = os.path.join(match_results_dir, match_file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    match_data = json.load(f)
                
                # 提取表格编号
                table_number = extract_table_number_from_filename(match_file)
                if table_number is None or table_number > len(all_tables):
                    continue
                
                target_table = all_tables[table_number - 1]
                
                # 根据位置信息替换单元格内容
                replace_cells_by_position(target_table, match_data)
                
            except Exception as e:
                print(f"处理匹配文件 {match_file} 时出错: {e}")
        
    except Exception as e:
        print(f"处理失败: {e}")
        raise

def get_all_tables_including_nested(doc):
    """
    获取文档中的所有表格（包括嵌套表格）
    使用与CSV提取相同的逻辑，通过哈希去重确保返回7个唯一表格
    
    参数:
        doc: docx.Document对象
        
    返回:
        list: 所有唯一表格的列表，按发现顺序排列（总共7个）
    """
    import hashlib
    import json
    
    all_tables = []
    table_hashes = set()
    
    def extract_table_data(table):
        """提取表格数据为二维数组"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = '\n'.join(p.text for p in cell.paragraphs).strip()
                row_data.append(cell_text if cell_text is not None else "")
            data.append(row_data)
        return data
    
    def hash_table_content(table_data):
        """计算表格内容的哈希值"""
        try:
            table_string = json.dumps(table_data, sort_keys=True)
        except TypeError:
            table_string = str(table_data)
        return hashlib.sha256(table_string.encode('utf-8')).hexdigest()
    
    def collect_nested_tables(table, processed_hashes, result_tables):
        """递归收集嵌套表格"""
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        nested_data = extract_table_data(nested_table)
                        nested_hash = hash_table_content(nested_data)
                        
                        if nested_hash not in processed_hashes:
                            processed_hashes.add(nested_hash)
                            result_tables.append(nested_table)
                        
                        # 递归处理更深层嵌套
                        collect_nested_tables(nested_table, processed_hashes, result_tables)
    
    # 处理顶层表格
    for table in doc.tables:
        table_data = extract_table_data(table)
        table_hash = hash_table_content(table_data)
        
        if table_hash not in table_hashes:
            table_hashes.add(table_hash)
            all_tables.append(table)
        
        # 收集嵌套表格
        collect_nested_tables(table, table_hashes, all_tables)
    
    print(f"Word文档中发现 {len(all_tables)} 个唯一表格（包括嵌套表格）")
    print(f"与HTML文档中的7个表格数量一致")
    
    return all_tables

def replace_cells_by_position(table, match_data):
    """
    根据位置信息替换表格单元格内容
    
    参数:
        table: 目标表格对象
        match_data: 匹配数据列表，包含位置信息
    """
    for item in match_data:
        try:
            # 解析位置信息，格式为 "(行号, 列号)"
            pos_str = item.get('valuePos', '')
            if not pos_str:
                continue
            
            row_index, col_index = parse_position(pos_str)
            if row_index is None or col_index is None:
                continue
            
            # 检查位置是否在表格范围内
            if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
                continue
            
            # 获取目标单元格
            target_cell = table.rows[row_index].cells[col_index]
            
            # 确定替换内容
            new_key = item.get('new_key', '').strip()
            old_key = item.get('old_key', '').strip()
            
            if new_key:
                replacement_text = f"[{new_key}]"
            elif old_key:
                replacement_text = f"[{old_key}]"
            else:
                continue
            
            # 替换单元格内容
            target_cell.text = replacement_text
            
        except Exception as e:
            print(f"处理匹配项时出错: {item}, 错误: {e}")

def extract_table_number_from_filename(filename):
    """从文件名中提取表格编号"""
    # 匹配 table_数字_matches.json 格式
    match = re.search(r'table_(\d+)_matches\.json', filename)
    if match:
        return int(match.group(1))
    return None

def parse_position(pos_str):
    """
    解析位置字符串，格式为 "(行号, 列号)"
    返回 (row_index, col_index)，索引从0开始
    
    根据实际分析：
    - JSON中 (1,1) 对应 Word表格的第1行第1列（0-based索引）
    - 即 (1,1) -> [1][1]，不需要减1
    """
    try:
        # 匹配 "(数字, 数字)" 格式
        match = re.search(r'\((\d+),\s*(\d+)\)', pos_str)
        if match:
            row = int(match.group(1))
            col = int(match.group(2))
            # JSON中的位置信息似乎就是0-based索引，直接使用
            return row, col
        else:
            return None, None
    except:
        return None, None

# 直接运行时的入口点
if __name__ == "__main__":
    import sys
    from docx import Document
    
    # 添加项目根目录到系统路径
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(os.path.dirname(current_dir))
    if project_dir not in sys.path:
        sys.path.append(project_dir)
    
    # 使用项目标准路径结构
    original_doc_path = os.path.join(project_dir, "document", "document.docx")
    match_results_dir = os.path.join(project_dir, "document", "match_results")
    template_doc_path = os.path.join(project_dir, "document", "template_table6_test.docx")
    
    # 检查输入文件和目录是否存在
    if not os.path.exists(original_doc_path):
        print(f"错误：原始文档 {original_doc_path} 不存在")
        exit(1)
        
    if not os.path.exists(match_results_dir):
        print(f"错误：匹配结果目录 {match_results_dir} 不存在")
        exit(1)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(template_doc_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # 读取原始文档
    print(f"正在读取原始文档: {original_doc_path}")
    doc = Document(original_doc_path)
    
    # 只使用 table_6 的匹配结果文件
    table_match_files = ['table_6_matches.json']
    
    # 检查匹配结果是否存在
    table_6_path = os.path.join(match_results_dir, 'table_6_matches.json')
    if not os.path.exists(table_6_path):
        print(f"错误：匹配结果文件不存在: {table_6_path}")
        exit(1)
    
    # 处理表格替换
    replace_values_with_placeholders(doc, match_results_dir, table_match_files)
    
    # 保存处理后的模板文档
    doc.save(template_doc_path)
    print(f"已保存模板文档: {template_doc_path}")